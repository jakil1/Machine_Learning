from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_homework_docx():
    doc = Document()

    # --- 标题 ---
    title = doc.add_heading('神经网络损失函数详解（启发式教学作业）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- 简介 ---
    p = doc.add_paragraph(
        '本次作业总结了深度学习中常用的7种损失函数，涵盖回归、分类及目标检测任务，并结合RT-DETR研究方向进行了举例说明。')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- 辅助函数：添加章节 ---
    def add_section(heading_text, content_list):
        doc.add_heading(heading_text, level=1)
        for item in content_list:
            # 损失函数名称
            h2 = doc.add_heading(item['name'], level=2)
            # 描述
            for desc in item['details']:
                doc.add_paragraph(desc, style='List Bullet')
            # 代码块
            if 'code' in item:
                p_code = doc.add_paragraph()
                runner = p_code.add_run(item['code'])
                runner.font.name = 'Courier New'  # 设置代码字体
                runner.font.size = Pt(9)
                runner.font.color.rgb = RGBColor(0, 50, 150)  # 蓝色高亮

    # --- 数据内容 ---

    # 1. 回归任务
    regression_data = [
        {
            'name': '1. 均方误差损失 (MSE Loss / L2 Loss)',
            'details': [
                '公式: L = Σ(y - ŷ)² / N',
                '特点: 对误差平方惩罚，误差越大惩罚越重。',
                '适用任务: 通用回归预测（如房价、坐标预测）。',
                'Pytorch模块:'
            ],
            'code': "import torch.nn as nn\ncriterion = nn.MSELoss()"
        },
        {
            'name': '2. 平均绝对误差损失 (L1 Loss)',
            'details': [
                '公式: L = Σ|y - ŷ| / N',
                '特点: 梯度恒定，对异常值不敏感。',
                '适用任务: 图像生成或对异常值敏感的回归。',
                'Pytorch模块:'
            ],
            'code': "import torch.nn as nn\ncriterion = nn.L1Loss()"
        },
        {
            'name': '3. 平滑 L1 损失 (Smooth L1 Loss)',
            'details': [
                '公式: 误差小用0.5x²，误差大用|x|-0.5',
                '特点: 结合了L1和L2的优点，收敛更稳定。',
                '适用任务: 目标检测中的边界框回归（如Faster R-CNN, SSD）。',
                'Pytorch模块:'
            ],
            'code': "import torch.nn as nn\ncriterion = nn.SmoothL1Loss()"
        }
    ]

    # 2. 分类任务
    classification_data = [
        {
            'name': '4. 交叉熵损失 (Cross Entropy Loss)',
            'details': [
                '公式: L = -Σ y * log(p)',
                '特点: 衡量两个概率分布的距离。',
                '适用任务: 多分类任务（如ImageNet分类，牛体部位分类）。',
                'Pytorch模块:'
            ],
            'code': "import torch.nn as nn\n# 输入不需要Softmax，函数内部已包含\ncriterion = nn.CrossEntropyLoss()"
        },
        {
            'name': '5. 二元交叉熵损失 (BCE Loss)',
            'details': [
                '公式: L = -[y*log(ŷ) + (1-y)*log(1-ŷ)]',
                '适用任务: 二分类或多标签分类（如判断图片是否有牛）。',
                'Pytorch模块:'
            ],
            'code': "import torch.nn as nn\ncriterion = nn.BCEWithLogitsLoss() # 推荐使用带Sigmoid的版本"
        }
    ]

    # 3. 目标检测（进阶）
    detection_data = [
        {
            'name': '6. Focal Loss (焦点损失)',
            'details': [
                '公式: FL = -α(1-p)^γ * log(p)',
                '特点: 增加权重因子，解决正负样本极度不平衡问题。',
                '适用任务: 单阶段目标检测（如RT-DETR, YOLO）的分类头。',
                'Pytorch模块:'
            ],
            'code': "from torchvision.ops import sigmoid_focal_loss"
        },
        {
            'name': '7. GIoU Loss (Generalized IoU Loss)',
            'details': [
                '公式: 考虑了预测框和真实框的重叠度及非重叠区域。',
                '特点: 解决了两框不相交时梯度为0的问题，比MSE更符合视觉直觉。',
                '适用任务: 现代目标检测的边界框回归（RT-DETR核心组件）。',
                'Pytorch模块:'
            ],
            'code': "from torchvision.ops import generalized_box_iou_loss"
        }
    ]

    # --- 生成文档 ---
    add_section('一、回归任务常用损失函数 (Regression)', regression_data)
    add_section('二、分类任务常用损失函数 (Classification)', classification_data)
    add_section('三、目标检测进阶损失函数 (与RT-DETR相关)', detection_data)

    # --- 结尾思考 ---
    doc.add_heading('四、启发式思考：RT-DETR中的实际应用', level=1)
    summary = doc.add_paragraph(
        '在我的牛体检测研究中，RT-DETR 模型实际上采用了组合损失函数策略：\n'
        '1. 使用 Focal Loss（BCE的改进）来处理分类任务，解决背景框过多的问题。\n'
        '2. 使用 L1 Loss 和 GIoU Loss 的加权组合来处理边界框回归，确保定位精准。'
    )
    summary.style = 'Intense Quote'

    # 保存
    save_name = '神经网络损失函数作业.docx'
    doc.save(save_name)
    print(f"✅ 成功生成Word文档：{save_name}")


if __name__ == '__main__':
    create_homework_docx()